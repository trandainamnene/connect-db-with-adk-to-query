"""
Script để xử lý file Word (IOS.docx và Android.docx):
1. Tách hình ảnh và lưu vào thư mục tương ứng
2. Parse nội dung thành JSON với đường dẫn hình ảnh
3. Lưu JSON để agent sử dụng
"""
import os
import json
import logging
import base64
import io
import re
import sys
import xml.etree.ElementTree as ET
from pathlib import Path
import pandas as pd

# Kiểm tra và import dependencies - không exit khi import, chỉ raise exception khi function được gọi
_import_errors = []

try:
    from PIL import Image
except ImportError as e:
    _import_errors.append("Pillow is not installed! Please run: pip install Pillow")
    Image = None

try:
    from unstructured.partition.docx import partition_docx
except ImportError as e:
    _import_errors.append("unstructured is not installed! Please run: pip install unstructured[docx]")
    partition_docx = None

try:
    from docx import Document
except ImportError as e:
    _import_errors.append("python-docx is not installed! Please run: pip install python-docx")
    Document = None

logging.basicConfig(level=logging.INFO)


def extract_images_from_docx(docx_path: str, output_folder: str) -> dict:
    """
    Tách hình ảnh từ file Word và lưu vào thư mục theo thứ tự xuất hiện trong document.
    Sử dụng python-docx để extract images theo thứ tự từ trên xuống dưới.
    
    Args:
        docx_path: Đường dẫn đến file .docx
        output_folder: Thư mục để lưu hình ảnh
    
    Returns:
        dict: Mapping giữa image index và file path
    """
    os.makedirs(output_folder, exist_ok=True)
    
    image_mapping = {}
    image_counter = 1
    seen_image_ids = set()  # Để tránh lưu lại ảnh đã lưu
    
    try:
        doc = Document(docx_path)
        
        # Extract images theo thứ tự xuất hiện trong document
        # Duyệt qua tất cả paragraphs và runs để tìm images
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                # Kiểm tra xem run có chứa image không
                if run._element.xml:
                    # Tìm tất cả image relationships trong run
                    try:
                        root = ET.fromstring(run._element.xml)
                        # Tìm tất cả elements có namespace drawing
                        for drawing in root.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                            r_embed = drawing.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if r_embed:
                                # Lấy image từ relationship
                                if r_embed in doc.part.rels:
                                    rel = doc.part.rels[r_embed]
                                    if "image" in rel.target_ref and r_embed not in seen_image_ids:
                                        try:
                                            image_data = rel.target_part.blob
                                            image = Image.open(io.BytesIO(image_data))
                                            
                                            # Lưu image
                                            image_filename = f"{image_counter}.jpg"
                                            image_path = os.path.join(output_folder, image_filename)
                                            
                                            # Convert to RGB nếu cần
                                            if image.mode in ('RGBA', 'LA', 'P'):
                                                background = Image.new('RGB', image.size, (255, 255, 255))
                                                if image.mode == 'P':
                                                    image = image.convert('RGBA')
                                                background.paste(image, mask=image.split()[-1] if image.mode in ('RGBA', 'LA') else None)
                                                image = background
                                            elif image.mode != 'RGB':
                                                image = image.convert('RGB')
                                            
                                            image.save(image_path, "JPEG", quality=85)
                                            image_mapping[image_counter] = image_path
                                            seen_image_ids.add(r_embed)
                                            logging.info(f"Saved image {image_counter} to {image_path}")
                                            image_counter += 1
                                            
                                        except Exception as e:
                                            logging.error(f"Error saving image {image_counter}: {e}")
                                            continue
                    except Exception as e:
                        # Nếu không parse được XML, bỏ qua
                        continue
        
        # Nếu không tìm thấy images bằng cách trên, fallback về cách cũ
        if len(image_mapping) == 0:
            logging.warning("No images found by parsing runs, trying relationships method...")
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    rel_id = rel.rId
                    if rel_id not in seen_image_ids:
                        try:
                            image_data = rel.target_part.blob
                            image = Image.open(io.BytesIO(image_data))
                            
                            # Lưu image
                            image_filename = f"{image_counter}.jpg"
                            image_path = os.path.join(output_folder, image_filename)
                            
                            # Convert to RGB nếu cần
                            if image.mode in ('RGBA', 'LA', 'P'):
                                background = Image.new('RGB', image.size, (255, 255, 255))
                                if image.mode == 'P':
                                    image = image.convert('RGBA')
                                background.paste(image, mask=image.split()[-1] if image.mode in ('RGBA', 'LA') else None)
                                image = background
                            elif image.mode != 'RGB':
                                image = image.convert('RGB')
                            
                            image.save(image_path, "JPEG", quality=85)
                            image_mapping[image_counter] = image_path
                            seen_image_ids.add(rel_id)
                            logging.info(f"Saved image {image_counter} to {image_path}")
                            image_counter += 1
                            
                        except Exception as e:
                            logging.error(f"Error saving image {image_counter}: {e}")
                            continue
                    
    except Exception as e:
        logging.error(f"Error extracting images from {docx_path}: {e}")
        import traceback
        logging.error(traceback.format_exc())
    
    logging.info(f"Extracted {len(image_mapping)} images in order")
    return image_mapping


def parse_docx_to_json(docx_path: str, image_mapping: dict, folder_type: str) -> list:
    """
    Parse nội dung Word thành JSON với đường dẫn hình ảnh.
    Sử dụng unstructured để parse text, kết hợp với image mapping.
    
    Args:
        docx_path: Đường dẫn đến file .docx
        image_mapping: Mapping giữa image index và file path
        folder_type: "IOS" hoặc "Android"
    
    Returns:
        list: Danh sách các step với text và image paths
    """
    # Partition docx file để lấy text structure
    elements = partition_docx(filename=docx_path)
    
    steps = []
    current_step = None
    step_number = 1
    image_index = 1
    
    logging.info(f"Parsing {len(elements)} elements from {docx_path}")
    
    for i, element in enumerate(elements):
        element_text = str(element).strip()
        
        if not element_text:
            continue
        
        # Kiểm tra nếu là step mới
        is_new_step = False
        step_match = re.search(r'Bước\s*(\d+)', element_text, re.IGNORECASE)
        if step_match:
            is_new_step = True
            step_number = int(step_match.group(1))
            logging.info(f"Found step {step_number}: {element_text[:50]}...")
        elif re.match(r'^\d+[\.\)]\s+', element_text):
            # Pattern: "1. " hoặc "1) "
            is_new_step = True
            step_number = int(re.match(r'^(\d+)', element_text).group(1))
            logging.info(f"Found step {step_number} (numbered): {element_text[:50]}...")
        elif re.match(r'^Step\s+(\d+)', element_text, re.IGNORECASE):
            # Pattern: "Step 1" hoặc "STEP 1"
            is_new_step = True
            step_number = int(re.match(r'^Step\s+(\d+)', element_text, re.IGNORECASE).group(1))
            logging.info(f"Found step {step_number} (Step format): {element_text[:50]}...")
        
        if is_new_step:
            # Lưu step trước đó
            if current_step:
                steps.append(current_step)
            
            # Kiểm tra xem element_text có chứa nhiều bước ngăn cách bằng "→", "->", hoặc ">" không
            # Android: "Bước 1: ... → Bước 2: ... → Bước 3: ..."
            # IOS: "Cài đặt > Quyền riêng tư > Dịch vụ định vị"
            if '→' in element_text or '->' in element_text:
                # Tách text thành các phần bằng arrow (Android)
                parts = re.split(r'[→]|->', element_text)
                for part_idx, part in enumerate(parts):
                    part = part.strip()
                    if part:
                        # Làm sạch text
                        part = re.sub(r'\s+', ' ', part).strip()
                        if part:
                            # Tìm step number trong phần này
                            step_match = re.search(r'Bước\s*(\d+)', part, re.IGNORECASE)
                            if step_match:
                                step_num = int(step_match.group(1))
                            else:
                                step_num = step_number + part_idx
                            
                            # Map image dựa trên step_number thực tế, không phải image_index tuần tự
                            # image_mapping có key là 1, 2, 3... tương ứng với 1.jpg, 2.jpg, 3.jpg...
                            image_path = image_mapping.get(step_num) if step_num in image_mapping else None
                            steps.append({
                                "step_number": step_num,
                                "text": part,
                                "image_path": image_path if image_path else None,
                                "folder_type": folder_type
                            })
                            # Cập nhật image_index để đảm bảo không bỏ sót
                            if step_num >= image_index:
                                image_index = step_num + 1
                # Không tạo current_step nữa vì đã tách thành nhiều steps
                current_step = None
            elif '>' in element_text and (folder_type == "IOS" or 'Bước' in element_text):
                # Tách text thành các phần bằng ">" (IOS hoặc có chứa "Bước")
                parts = re.split(r'\s*>\s*', element_text)
                for part_idx, part in enumerate(parts):
                    part = part.strip()
                    if part:
                        # Làm sạch text
                        part = re.sub(r'\s+', ' ', part).strip()
                        if part:
                            # Tìm step number trong phần này (có thể có "Bước X:" hoặc không)
                            step_match = re.search(r'Bước\s*(\d+)', part, re.IGNORECASE)
                            if step_match:
                                step_num = int(step_match.group(1))
                            else:
                                # Nếu không có "Bước X:", thêm vào đầu
                                step_num = step_number + part_idx
                                # Format: "Bước X: [text]" nếu chưa có
                                if not part.startswith('Bước'):
                                    part = f"Bước {step_num}: {part}"
                            
                            # Map image dựa trên step_number thực tế, không phải image_index tuần tự
                            image_path = image_mapping.get(step_num) if step_num in image_mapping else None
                            steps.append({
                                "step_number": step_num,
                                "text": part,
                                "image_path": image_path if image_path else None,
                                "folder_type": folder_type
                            })
                            # Cập nhật image_index để đảm bảo không bỏ sót
                            if step_num >= image_index:
                                image_index = step_num + 1
                # Không tạo current_step nữa vì đã tách thành nhiều steps
                current_step = None
            else:
                # Tạo step mới với image nếu có
                image_path = image_mapping.get(image_index)
                current_step = {
                    "step_number": step_number,
                    "text": element_text,
                    "image_path": image_path if image_path else None,
                    "folder_type": folder_type
                }
                
                if image_path:
                    image_index += 1
        else:
            # Thêm text vào step hiện tại
            if current_step:
                if current_step["text"]:
                    current_step["text"] += "\n" + element_text
                else:
                    current_step["text"] = element_text
            else:
                # Nếu chưa có step nào và có text, tạo step đầu tiên
                # (có thể document không có format "Bước X")
                if element_text and len(element_text) > 10:  # Chỉ tạo step nếu text đủ dài
                    image_path = image_mapping.get(image_index)
                    current_step = {
                        "step_number": step_number,
                        "text": element_text,
                        "image_path": image_path if image_path else None,
                        "folder_type": folder_type
                    }
                    if image_path:
                        image_index += 1
    
    # Thêm step cuối cùng
    if current_step:
        steps.append(current_step)
    
    logging.info(f"Parsed {len(steps)} steps from {docx_path}")
    
    if len(steps) == 0:
        logging.warning("No steps found with pattern matching, trying alternative parsing...")
        
        steps = _parse_by_paragraphs(docx_path, image_mapping, folder_type)
        
        if len(steps) == 0:
            if len(image_mapping) == 0:
                current_dir = os.path.dirname(os.path.abspath(docx_path))
                image_folder = os.path.join(current_dir, f"{folder_type}_Instruction")
                
                if os.path.exists(image_folder):
                    img_files = sorted([f for f in os.listdir(image_folder) 
                                      if f.lower().endswith(('.jpg', '.jpeg', '.png'))])
                    for i, img_file in enumerate(img_files, 1):
                        image_mapping[i] = os.path.join(image_folder, img_file)
            
            if len(image_mapping) > 0:
                logging.warning(f"No text steps found, creating steps from {len(image_mapping)} images...")
                try:
                    steps = _create_steps_from_images(docx_path, image_mapping, folder_type)
                except Exception as e:
                    logging.error(f"Error creating steps: {e}")
                    import traceback
                    traceback.print_exc()
    return steps


def _parse_by_paragraphs(docx_path: str, image_mapping: dict, folder_type: str) -> list:
    """
    Parse bằng cách chia theo paragraphs nếu không tìm thấy pattern "Bước X".
    """
    doc = Document(docx_path)
    steps = []
    step_number = 1
    image_index = 1
    current_text = []
    
    # Lấy tất cả paragraphs có text
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    if not paragraphs:
        return steps
    
    # Chia paragraphs thành các nhóm (steps) dựa vào số lượng images
    num_images = len(image_mapping)
    if num_images > 0:
        # Chia đều paragraphs cho các steps
        paragraphs_per_step = max(1, len(paragraphs) // num_images)
    else:
        paragraphs_per_step = 3  # Mặc định 3 paragraphs mỗi step
    
    for i, text in enumerate(paragraphs):
        current_text.append(text)
        
        # Tạo step mới khi đủ số paragraphs hoặc hết paragraphs
        should_create_step = (
            len(current_text) >= paragraphs_per_step or 
            i == len(paragraphs) - 1 or
            (num_images > 0 and len(steps) < num_images and len(current_text) >= 2)
        )
        
        if should_create_step and current_text:
            image_path = image_mapping.get(image_index) if image_index <= len(image_mapping) else None
            steps.append({
                "step_number": step_number,
                "text": "\n".join(current_text),
                "image_path": image_path if image_path else None,
                "folder_type": folder_type
            })
            if image_path:
                image_index += 1
            step_number += 1
            current_text = []
    
    logging.info(f"Parsed {len(steps)} steps using paragraph-based method")
    return steps


def _get_guide_from_excel(folder_type: str) -> str:
    """Lấy hướng dẫn từ Excel dựa trên folder_type."""
    try:
        current_dir = os.path.dirname(os.path.abspath(__file__))
        excel_path = os.path.join(current_dir, "device_models_with_location.xlsx")
        
        if not os.path.exists(excel_path):
            logging.warning(f"Excel file not found: {excel_path}")
            return ""
        
        df = pd.read_excel(excel_path).dropna(subset=['How_to_Enable_Location'])
        
        if df.empty:
            return ""
        
        if folder_type == "IOS":
            mask = df['ModelName'].astype(str).str.contains('iPhone|iOS|iPad', case=False, na=False, regex=True)
        else:
            mask = ~df['ModelName'].astype(str).str.contains('iPhone|iOS|iPad', case=False, na=False, regex=True)
        
        matches = df[mask]
        if not matches.empty:
            result = matches.iloc[0]
            guide_text_raw = result.get('How_to_Enable_Location', '')
            guide_text = "" if pd.isna(guide_text_raw) else str(guide_text_raw).strip()
            
            if guide_text:
                model_name = str(result.get('ModelName', ''))
                logging.info(f"Found guide from Excel for {model_name} ({folder_type})")
                return guide_text
            
    except Exception as e:
        logging.error(f"Error reading Excel: {e}")
        import traceback
        logging.error(traceback.format_exc())
    
    return ""


def _create_steps_from_images(docx_path: str, image_mapping: dict, folder_type: str) -> list:
    """
    Tạo steps dựa vào số lượng images và text từ Excel.
    Mỗi image sẽ là một step với text từ Excel.
    """
    num_images = len(image_mapping)
    if num_images == 0:
        return []
    
    guide_text = _get_guide_from_excel(folder_type)
    text_parts = []
    
    if guide_text and len(guide_text.strip()) > 10:
        if '→' in guide_text or '->' in guide_text:
            parts = re.split(r'[→]|->', guide_text)
            for part in parts:
                part = re.sub(r'\s+', ' ', part.strip())
                if part:
                    text_parts.append(part)
        elif '>' in guide_text and ('Bước' in guide_text or folder_type == "IOS"):
            parts = re.split(r'\s*>\s*', guide_text)
            for part in parts:
                part = re.sub(r'\s+', ' ', part.strip())
                if part:
                    text_parts.append(part)
        
        if len(text_parts) == 0:
            lines = [l.strip() for l in guide_text.split('\n') if l.strip()]
            if lines:
                current_step_text = []
                for line in lines:
                    if re.search(r'Bước\s*(\d+)', line, re.IGNORECASE) or re.match(r'^(\d+)[\.\)]\s+', line):
                        if current_step_text:
                            text_parts.append('\n'.join(current_step_text))
                            current_step_text = []
                        current_step_text.append(line)
                    else:
                        current_step_text.append(line)
                if current_step_text:
                    text_parts.append('\n'.join(current_step_text))
            
            if len(text_parts) == 0:
                lines_per_step = max(1, len(lines) // num_images) if num_images > 0 else len(lines)
                for i in range(0, len(lines), lines_per_step):
                    part = '\n'.join(lines[i:i+lines_per_step])
                    if part:
                        text_parts.append(part)
        
        if len(text_parts) < num_images:
            while len(text_parts) < num_images:
                text_parts.append(text_parts[-1] if text_parts else "Hướng dẫn")
        elif len(text_parts) > num_images:
            text_parts = text_parts[:num_images]
    else:
        if guide_text and len(guide_text.strip()) > 0:
            lines = [l.strip() for l in guide_text.split('\n') if l.strip()]
            if lines:
                lines_per_step = max(1, len(lines) // num_images) if num_images > 0 else len(lines)
                for i in range(0, len(lines), lines_per_step):
                    part = '\n'.join(lines[i:i+lines_per_step])
                    if part:
                        text_parts.append(part)
                while len(text_parts) < num_images:
                    text_parts.append(text_parts[-1] if text_parts else guide_text[:50])
                text_parts = text_parts[:num_images]
        
        if len(text_parts) == 0:
            text_parts = [f"Hướng dẫn bước {i + 1}" for i in range(num_images)]
    
    steps = []
    current_dir = os.path.dirname(os.path.abspath(docx_path))
    
    for i, image_key in enumerate(sorted(image_mapping.keys()), 1):
        image_path = image_mapping[image_key]
        
        if not os.path.isabs(image_path):
            image_path = os.path.join(current_dir, image_path)
        
        if not os.path.exists(image_path):
            logging.warning(f"Image path does not exist: {image_path}")
            text = text_parts[i-1] if i <= len(text_parts) else f"Hướng dẫn bước {i}"
            steps.append({
                "step_number": image_key,
                "text": text,
                "image_path": None,
                "folder_type": folder_type
            })
            continue
        
        if i <= len(text_parts) and text_parts[i-1]:
            text = text_parts[i-1]
            if text == f"Hướng dẫn bước {i}" and guide_text:
                lines = [l.strip() for l in guide_text.split('\n') if l.strip()]
                if lines:
                    lines_per_step = max(1, len(lines) // num_images)
                    start_idx = (i - 1) * lines_per_step
                    end_idx = start_idx + lines_per_step if i < num_images else len(lines)
                    text = '\n'.join(lines[start_idx:end_idx]) if start_idx < len(lines) else text
        else:
            text = f"Hướng dẫn bước {i}"
        
        step_match = re.search(r'Bước\s*(\d+)', text, re.IGNORECASE)
        if step_match and int(step_match.group(1)) != image_key:
            text = re.sub(r'Bước\s*\d+', f'Bước {image_key}', text, count=1, flags=re.IGNORECASE)
        
        steps.append({
            "step_number": image_key,
            "text": text,
            "image_path": image_path,
            "folder_type": folder_type
        })
    
    logging.info(f"Created {len(steps)} steps from {num_images} images")
    return steps


def process_docx_files():
    """Xử lý cả 2 file IOS.docx và Android.docx."""
    if _import_errors:
        python_exe = sys.executable
        in_venv = hasattr(sys, 'real_prefix') or (
            hasattr(sys, 'base_prefix') and sys.base_prefix != sys.prefix
        )
        
        error_msg = "Missing required dependencies:\n" + "\n".join(f"  - {err}" for err in _import_errors)
        env_info = f"Python: {python_exe}, In venv: {in_venv}"
        
        logging.error(f"{error_msg}\n{env_info}")
        raise ImportError(f"{error_msg}\n{env_info}\nCài đặt: pip install -r requirements.txt")
    
    if Image is None or partition_docx is None or Document is None:
        raise ImportError("Required dependencies are not available. Install: pip install -r requirements.txt")
    
    current_dir = os.path.dirname(os.path.abspath(__file__))
    logging.info("Starting Word file processing...")
    
    ios_docx = os.path.join(current_dir, "IOS.docx")
    ios_output_folder = os.path.join(current_dir, "IOS_Instruction")
    
    if os.path.exists(ios_docx):
        logging.info("Processing IOS.docx...")
        ios_images = extract_images_from_docx(ios_docx, ios_output_folder)
        
        if len(ios_images) == 0 and os.path.exists(ios_output_folder):
            existing_images = sorted([f for f in os.listdir(ios_output_folder) 
                                    if f.lower().endswith(('.jpg', '.jpeg', '.png'))])
            for i, img_file in enumerate(existing_images, 1):
                ios_images[i] = os.path.join(ios_output_folder, img_file)
        
        ios_steps = parse_docx_to_json(ios_docx, ios_images, "IOS")
        ios_json_path = os.path.join(current_dir, "ios_instructions.json")
        with open(ios_json_path, 'w', encoding='utf-8') as f:
            json.dump(ios_steps, f, ensure_ascii=False, indent=2)
        logging.info(f"Saved {len(ios_steps)} IOS steps to {ios_json_path}")
    else:
        logging.warning(f"IOS.docx not found at {ios_docx}")
    
    android_docx = os.path.join(current_dir, "Android.docx")
    android_output_folder = os.path.join(current_dir, "Android_Instruction")
    
    if os.path.exists(android_docx):
        logging.info("Processing Android.docx...")
        android_images = extract_images_from_docx(android_docx, android_output_folder)
        
        if len(android_images) == 0 and os.path.exists(android_output_folder):
            existing_images = sorted([f for f in os.listdir(android_output_folder) 
                                    if f.lower().endswith(('.jpg', '.jpeg', '.png'))])
            for i, img_file in enumerate(existing_images, 1):
                android_images[i] = os.path.join(android_output_folder, img_file)
        
        android_steps = parse_docx_to_json(android_docx, android_images, "Android")
        android_json_path = os.path.join(current_dir, "android_instructions.json")
        with open(android_json_path, 'w', encoding='utf-8') as f:
            json.dump(android_steps, f, ensure_ascii=False, indent=2)
        logging.info(f"Saved {len(android_steps)} Android steps to {android_json_path}")
    else:
        logging.warning(f"Android.docx not found at {android_docx}")
    
    logging.info("Word file processing completed")


if __name__ == "__main__":
    try:
        process_docx_files()
    except Exception as e:
        logging.error(f"Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
